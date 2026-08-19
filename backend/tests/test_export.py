"""Tests for exact-format paper export (docx template clone + layout styling)."""
import io
import os

import pytest

import api


def _sample_docx(path):
    """A realistic institute sample: header, footer, Times font, custom margin."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    sample = Document()
    sample.styles["Normal"].font.name = "Times New Roman"
    sample.styles["Normal"].font.size = Pt(12)
    sec = sample.sections[0]
    sec.left_margin = Inches(0.7)
    h = sec.header.paragraphs[0]
    h.text = "GOVT. COLLEGE OF SCIENCE - MID-TERM EXAMINATION"
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sec.footer.paragraphs[0].text = "Good Luck!"
    sample.add_paragraph("Q1. Old question that should be REPLACED. (10 marks)")
    sample.save(str(path))
    return str(path)


def test_docx_template_clone(tmp_path):
    """A .docx sample's header/footer/font/margins survive; body is replaced."""
    from docx import Document
    pattern = _sample_docx(tmp_path / "sample.docx")
    new_paper = "# MID-TERM 2026\n\n## Section A\n\nQ1. Define photosynthesis. (5 marks)"
    data, mime, ext = api._paper_to_bytes(new_paper, "docx", layout={}, pattern_path=pattern)
    out = Document(io.BytesIO(data))
    sec = out.sections[0]
    body = "\n".join(p.text for p in out.paragraphs)
    assert "GOVT. COLLEGE OF SCIENCE" in sec.header.paragraphs[0].text
    assert sec.footer.paragraphs[0].text == "Good Luck!"
    assert out.styles["Normal"].font.name == "Times New Roman"
    assert round(sec.left_margin.inches, 2) == 0.7
    assert "REPLACED" not in body
    assert "photosynthesis" in body


def test_layout_styled_export():
    """Photo-derived layout spec: serif font, centered title block, footer."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    paper = "MY ACADEMY\nTEST - BIOLOGY\n\nQ1. Define osmosis. (5 marks)"
    layout = {"font": "serif", "center_top_lines": 2, "footer": "My Academy"}
    data, _, _ = api._paper_to_bytes(paper, "docx", layout=layout)
    out = Document(io.BytesIO(data))
    assert out.styles["Normal"].font.name == "Times New Roman"
    assert out.sections[0].footer.paragraphs[0].text == "My Academy"
    first = [p for p in out.paragraphs if p.text.strip()][0]
    assert first.alignment == WD_ALIGN_PARAGRAPH.CENTER

    pdf, _, _ = api._paper_to_bytes(paper, "pdf", layout=layout)
    assert pdf[:4] == b"%PDF"
    assert "Times" in pdf.decode("latin-1", "ignore")


def test_pdf_typography_normalized_but_urdu_refused():
    """Em-dashes/smart quotes export fine; Urdu still gets the friendly refusal."""
    data, _, _ = api._paper_to_bytes("Q1 — define ‘osmosis’ … (5 marks)", "pdf")
    assert data[:4] == b"%PDF"
    with pytest.raises(ValueError):
        api._paper_to_bytes("سوال نمبر ایک", "pdf")
